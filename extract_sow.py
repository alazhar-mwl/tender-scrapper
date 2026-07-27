"""
Phase 3 — Scope-of-Work extractor (fully offline — never touches the portal)
-----------------------------------------------------------------------------
For each tender in tenders.json that has downloaded documents, extracts text
from the PDF/DOCX files in documents/<ref>/ (looking inside ZIPs too), finds
the "Scope of Work" section by heading detection, and writes these fields
back into tenders.json:

    scope_of_work   — the extracted section (or full-text fallback)
    sow_source      — which file it came from
    sow_extraction  — "heading-match" | "fallback-fulltext" | "failed" | "no-documents"

Setup:  pip install pymupdf python-docx
Run:    python extract_sow.py          # only tenders not yet extracted
        python extract_sow.py --all    # re-extract everything
"""

import json
import re
import sys
import zipfile
from pathlib import Path

# Windows consoles default to cp1252, which can't encode ✓/✗ — degrade, don't crash
if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
    sys.stdout.reconfigure(errors="replace")

BASE_DIR = Path(__file__).parent
OUT_FILE = BASE_DIR / "tenders.json"
DOCS_DIR = BASE_DIR / "documents"

MAX_SOW_CHARS      = 6_000   # cap on an extracted SOW section
FALLBACK_CHARS     = 3_000   # cap on the full-text fallback
MIN_SECTION_CHARS  = 80      # shorter than this ⇒ treat heading match as false positive

# "3. SCOPE OF WORK", "B) Scope of Services", "SCOPE OF WORK:", etc.
SOW_HEADING = re.compile(
    r"^[ \t]*(?:(?:\d+(?:\.\d+)*|[A-Z]|[IVXivx]+)[.)\s:–-]{1,4})?[ \t]*"
    r"(scope\s+of\s+works?|scope\s+of\s+services?|scope\s+of\s+supply|"
    r"work\s*scope|description\s+of\s+(?:the\s+)?works?)\b[^\n]{0,40}$",
    re.IGNORECASE | re.MULTILINE,
)
# A subsequent numbered/lettered heading in CAPS or Title Case ends the section
NEXT_HEADING = re.compile(
    r"^[ \t]*(?:\d+(?:\.\d+)*|[A-Z]|[IVX]+)[.)\s:–-]{1,4}[ \t]*[A-Z][A-Za-z ()/&,-]{3,60}$",
    re.MULTILINE,
)


def pdf_text(path: Path) -> str:
    import fitz  # PyMuPDF
    with fitz.open(str(path)) as doc:
        return "\n".join(page.get_text() for page in doc)


def docx_text(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("  ".join(c.text for c in row.cells))
    return "\n".join(parts)


def file_text(path: Path) -> str:
    """Extract text from one file; returns '' when the format is unsupported."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return pdf_text(path)
        if suffix == ".docx":
            return docx_text(path)
        if suffix in (".txt", ".csv"):
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".doc":
            print(f"    ! {path.name}: legacy .doc not supported — convert to .docx to extract")
    except Exception as exc:
        print(f"    ! {path.name}: extraction failed — {exc}")
    return ""


def unpack_zips(folder: Path) -> None:
    """Extract any ZIPs in the folder into subfolders (idempotent)."""
    for zpath in folder.glob("*.zip"):
        target = folder / f"{zpath.stem}_unzipped"
        if target.exists():
            continue
        try:
            with zipfile.ZipFile(zpath) as zf:
                zf.extractall(target)
            print(f"    unpacked {zpath.name}")
        except Exception as exc:
            print(f"    ! {zpath.name}: unzip failed — {exc}")


def find_sow(text: str) -> str | None:
    """Return the Scope of Work section, or None if no heading matches."""
    m = SOW_HEADING.search(text)
    if not m:
        return None
    start = m.start()
    # Look for the next heading a little past the SOW heading itself
    nxt = NEXT_HEADING.search(text, m.end() + 40)
    end = min(nxt.start() if nxt else len(text), start + MAX_SOW_CHARS)
    section = text[start:end].strip()
    return section if len(section) >= MIN_SECTION_CHARS else None


def parse_srm_notes(text: str) -> str | None:
    """
    Pull the Notes-table rows out of a saved SRM detail page (notes.txt).
    SAP's innerText renders each note as one tab-separated line:
        \\tDocument Header\\tClauses\\t1.0 Seller is invited to submit ...
        \\tItem01"O-RING,..."\\tPR Material Text\\tSpare_parts Part name ...
    Everything else on the page (toolbar labels, empty Attachments table,
    accessibility hints) is noise — keep only the note rows.
    """
    rows = []
    for line in text.splitlines():
        parts = [p.strip() for p in line.split("\t") if p.strip()]
        if len(parts) != 3:
            continue
        assigned, category, preview = parts
        cat = category.lower()
        if cat == "clauses" or cat.endswith("text"):
            rows.append(f"{assigned} — {category}: {preview}")
    if not rows:
        return None
    section = "\n".join(dict.fromkeys(rows))          # de-dup, keep order
    return section if len(section) >= MIN_SECTION_CHARS else None


def candidate_files(folder: Path) -> list[Path]:
    """All extractable files, SOW-named ones first, then PDFs, then the rest."""
    files = [p for p in folder.rglob("*")
             if p.is_file() and p.suffix.lower() in (".pdf", ".docx", ".doc", ".txt", ".csv")]

    def priority(p: Path) -> tuple:
        name = p.name.lower()
        named_sow = not ("sow" in name or "scope" in name)   # False sorts first
        return (named_sow, p.suffix.lower() != ".pdf", name)

    return sorted(files, key=priority)


def process_tender(tender: dict) -> None:
    ref = tender.get("reference_number", "")
    folder = DOCS_DIR / ref
    if not ref or not folder.is_dir() or not any(folder.iterdir()):
        tender["scope_of_work"] = ""
        tender["sow_source"] = ""
        tender["sow_extraction"] = "no-documents"
        return

    print(f"  {ref}:")
    unpack_zips(folder)

    # Files we ourselves saved from the portal's own detail-page text —
    # notes.txt (PDO) / description.txt (OQ). These are ALWAYS tender-specific
    # by construction (freshly scraped per tender), unlike downloaded
    # attachments which can be generic boilerplate (e.g. OQ's standard
    # "Terms & Conditions" PDF, identical across every tender) that happens
    # to have more raw text and would otherwise win a naive longest-text
    # fallback. Prefer these over any other fallback when no SOW heading matches.
    SCRAPED_TEXT_FILES = ("notes.txt", "description.txt")

    fallback_text, fallback_file = "", ""
    scraped_sow, scraped_source = None, ""
    for path in candidate_files(folder):
        text = file_text(path)
        if not text.strip():
            continue
        if len(text) > len(fallback_text):
            fallback_text, fallback_file = text, path.name
        section = find_sow(text)
        if section:
            tender["scope_of_work"] = section
            tender["sow_source"] = path.name
            tender["sow_extraction"] = "heading-match"
            print(f"    ✓ SOW found in {path.name} ({len(section)} chars)")
            return
        if path.name in SCRAPED_TEXT_FILES and scraped_sow is None:
            parsed = parse_srm_notes(text) if path.name == "notes.txt" else None
            candidate = parsed if parsed else (text.strip() if len(text.strip()) >= MIN_SECTION_CHARS else None)
            if candidate:
                scraped_sow, scraped_source = candidate, path.name

    if scraped_sow:
        tender["scope_of_work"] = scraped_sow[:MAX_SOW_CHARS]
        tender["sow_source"] = scraped_source
        tender["sow_extraction"] = "notes-table" if scraped_source == "notes.txt" else "page-text"
        print(f"    ✓ SOW built from {scraped_source} ({len(scraped_sow)} chars)")
    elif fallback_text:
        tender["scope_of_work"] = fallback_text[:FALLBACK_CHARS].strip()
        tender["sow_source"] = fallback_file
        tender["sow_extraction"] = "fallback-fulltext"
        print(f"    ~ no SOW heading — stored start of {fallback_file}")
    else:
        tender["scope_of_work"] = ""
        tender["sow_source"] = ""
        tender["sow_extraction"] = "failed"
        print("    ✗ no text could be extracted from any document")


def main() -> None:
    if not OUT_FILE.exists():
        sys.exit("tenders.json not found — run tender_scraper.py first.")
    tenders = json.loads(OUT_FILE.read_text(encoding="utf-8"))
    redo_all = "--all" in sys.argv

    todo = [t for t in tenders
            if redo_all or not t.get("sow_extraction")
            or t.get("sow_extraction") in ("failed", "no-documents")]
    print(f"Extracting scope of work for {len(todo)} tender(s)…")

    for tender in todo:
        process_tender(tender)

    OUT_FILE.write_text(
        json.dumps(tenders, indent=2, ensure_ascii=False, default=str),
        encoding="utf-8",
    )
    counts: dict[str, int] = {}
    for t in tenders:
        key = t.get("sow_extraction", "not-run")
        counts[key] = counts.get(key, 0) + 1
    print("Done. Results:", ", ".join(f"{k}={v}" for k, v in sorted(counts.items())))


if __name__ == "__main__":
    main()
